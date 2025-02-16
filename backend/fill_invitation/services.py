from django.core.files.uploadedfile import InMemoryUploadedFile

from openai import OpenAI
from typing import Generator
from dotenv.main import load_dotenv
from docx import Document
from docx.shared import Pt

import base64
import json
import os

load_dotenv()


def get_base64_image(images: list[InMemoryUploadedFile] | InMemoryUploadedFile) -> Generator[str, None, None]:
    """Yields base64-encoded images from uploaded files in the request."""
    if isinstance(images, InMemoryUploadedFile):
        images = [images]

    try:
        for image in images:
            base64_image = base64.b64encode(image.read()).decode('utf-8')
            image.seek(0)
            yield base64_image
    except Exception as e:
        raise Exception(str(e))


def retrieve_from_image(images: list[InMemoryUploadedFile]) -> list[dict[str, str]]:
    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

        passport_results = []

        for indx, base64_image in enumerate(get_base64_image(images)):
            
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": [
                            {
                                "type": "text",
                                "text": "You are a passport OCR assistant. Extract only the following "
                                        "details from the image and return the output strictly as a JSON"
                                        " object with the following structure:"
                                        "\n{\n"
                                        "  \"surname\": \"\",\n"
                                        "  \"name\": \"\",\n"
                                        "  \"father_name\": \"\",\n"
                                        "  \"date_of_issue\": \"YYYY-MM-DD\",\n"
                                        "  \"passport_number\": \"\",\n"
                                        "  \"date_of_birth\": \"YYYY-MM-DD\"\n"
                                        "}"
                            }
                        ],
                    },
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "Extract passport details from this image:"},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                        ]
                    }
                ],
                max_tokens=300
            )

            gpt_response = response.choices[0].message.content.strip()
            if not gpt_response:
                raise TypeError("Empty response from OpenAI")

            if gpt_response.startswith("```json") and gpt_response.endswith("```"):
                gpt_response = gpt_response[7:-3].strip()
        
            passport_data = json.loads(gpt_response)
            passport_data['id'] = indx + 1
            passport_data['surname'] = passport_data['surname'].upper()
            passport_data['name'] = passport_data['name'].upper()
            passport_data['father_name'] = passport_data['father_name'].upper()

            passport_results.append(passport_data)

    except Exception as e:
            raise Exception(str(e))

    return passport_results


def is_foreign_passport(image: InMemoryUploadedFile) -> bool:

    try:
        client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

        for base64_image in get_base64_image(image):
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an AI assistant that determines if an image contains a foreign passport. "
                                    "Return only one of the following JSON responses strictly: {\"result\": true} "
                                    "if it is a foreign passport, or {\"result\": false} otherwise."
                    },
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "Does this image contain a foreign passport?"},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                        ]
                    }
                ],
                max_tokens=20
            )

            gpt_response = response.choices[0].message.content.strip()
            if not gpt_response:
                raise TypeError("Empty response from OpenAI")

         
            result_data = json.loads(gpt_response)
            return result_data.get("result", False)
        
    except Exception as e:
        raise Exception(str(e))
    

def translate_text(text: str) -> str:
    """Ensures OpenAI correctly translates names from Uzbek, Kazakh, and other Asian languages into Ukrainian."""
    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a professional document translator specializing in passports and official documents."
                    " Translate name surname and father`s name into Ukrainian."
                    " Always ensure:"
                    " - 'KIZI' is translated as 'КІЗІ'."
                    " - Do not add explanations, just return the translated text."
                ),
            },
            {
                "role": "user",
                "content": f"Translate this name into Ukrainian: {text}"
            }
        ],
        max_tokens=100
    )

    return response.choices[0].message.content.strip().upper()


def translate_passport_list(passport_list: list[dict]) -> list[dict]:
    """Iterates over the list and translates `surname`, `name`, and `father_name` fields."""
    translated_passports = []

    for passport in passport_list:
        translated_passport = passport.copy()
        translated_passport["surname"] = translate_text(passport["surname"])
        translated_passport["name"] = translate_text(passport["name"])
        translated_passport["father_name"] = translate_text(passport["father_name"])
        translated_passports.append(translated_passport)

    return translated_passports


def merge_passport_data(data_list: list, data_ukr_list: list) -> list[dict]:
    """ Merges two lists of passport data (English & Ukrainian) into a list of dictionaries. """
    
    merged_data_list = [] 
    
    # Convert lists into dictionaries indexed by ID for quick lookup
    data_dict = {entry["id"]: entry for entry in data_list}
    data_ukr_dict = {entry["id"]: entry for entry in data_ukr_list}

    # Process each matching entry
    for id_value, eng_entry in data_dict.items():
        if id_value in data_ukr_dict:
            ukr_entry = data_ukr_dict[id_value]  # Get corresponding Ukrainian entry

            # Create a merged dictionary for this ID
            merged_entry = {
                "[SNF]": f'{eng_entry["surname"]} {eng_entry["name"]} {eng_entry["father_name"]}',
                "[ФИО]": f'{ukr_entry["surname"]} {ukr_entry["name"]} {ukr_entry["father_name"]}',

                "[Birth date]": eng_entry["date_of_birth"],
                "[Дата рождения]": ukr_entry["date_of_birth"],

                "[Passport number]": eng_entry["passport_number"],
                "[Номер паспорта]": ukr_entry["passport_number"],

                "[Date of issue]": eng_entry["date_of_issue"],
                "[Дата выдачи]": ukr_entry["date_of_issue"],
            }
            
            merged_data_list.append(merged_entry) 

    return merged_data_list 


def replace_text_in_runs(paragraph, replacements):
    """ Replaces placeholders inside paragraph runs while preserving formatting. """
    
    # Define which placeholders should be bold after replacement
    BOLD_FIELDS = ["[ФИО]", "[SNF]"]
    
    for key, val in replacements.items():
        for run in paragraph.runs:
            if key in run.text:
                # Preserve formatting from the first occurrence of the key
                font_style = {
                    "bold": run.bold or (key in BOLD_FIELDS),  # Ensure bold if in BOLD_FIELDS
                    "italic": run.italic,
                    "underline": run.underline,
                    "size": run.font.size if run.font.size else Pt(11),  # Default size
                }

                # Replace text while keeping formatting
                run.text = run.text.replace(key, val)
                run.bold = font_style["bold"]
                run.italic = font_style["italic"]
                run.underline = font_style["underline"]
                run.font.size = font_style["size"]
                

def doc(data_list: list[dict]) -> None:
    """ Generates multiple Word documents, replacing placeholders for each person while preserving formatting. """

    template_path = os.path.abspath("./fill_invitation/invitation_docx/invitation.docx")

    for idx, data in enumerate(data_list, start=1):
        doc = Document(template_path)

        # Iterate through all paragraphs and replace text inside runs
        for paragraph in doc.paragraphs:
            replace_text_in_runs(paragraph, data)

        # Iterate through tables (if placeholders are inside tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_runs(paragraph, data)

        output_filename = f"/home/rodion/Public/filled_{idx}.docx"
        doc.save(output_filename)